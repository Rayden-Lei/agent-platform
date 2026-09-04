"""把 docs/01～04 四份 Markdown 设计文档渲染成一份 Word 文档 docs/agent-platform-design.docx。

用法：在仓库根目录执行 backend/.venv/bin/python gen_docx.py
Markdown 是唯一的源，docx 是生成产物：改了源文档就重新生成并一起提交，不要手改 docx。

支持的 Markdown 子集：# 标题（1～3 级）、段落、- 列表、1. 列表、| 表格 |、``` 代码块、> 引用、
**粗体**、`行内代码`、[文字](链接)。其他语法按普通文本输出。
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SOURCES = ["01-需求说明.md", "02-架构设计.md", "03-数据库设计.md", "04-接口设计.md"]
OUTPUT = ROOT / "docs" / "agent-platform-design.docx"

FONT_BODY, FONT_HEAD, FONT_MONO = "宋体", "黑体", "Consolas"
HEADING_SIZES = {1: 16, 2: 13, 3: 12}

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[([^\]]+)\]\([^)]+\))")
TABLE_SEP = re.compile(r"^\|?\s*:?-{2,}")


def _font(run, name, size, bold=False, italic=False, mono=False):
    run.font.name = FONT_MONO if mono else name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MONO if mono else name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def _inline(par, text, size=10.5, bold=False, italic=False, font=FONT_BODY):
    """渲染行内粗体 / 行内代码 / 链接，其余原样输出。"""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            _font(par.add_run(text[pos:m.start()]), font, size, bold, italic)
        tok = m.group(0)
        if tok.startswith("**"):
            _font(par.add_run(tok[2:-2]), font, size, True, italic)
        elif tok.startswith("`"):
            _font(par.add_run(tok[1:-1]), font, size - 0.5, bold, italic, mono=True)
        else:
            _font(par.add_run(m.group(2)), font, size, bold, italic)
        pos = m.end()
    if pos < len(text):
        _font(par.add_run(text[pos:]), font, size, bold, italic)


class Renderer:
    def __init__(self, doc: Document):
        self.doc = doc

    def heading(self, level: int, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
        p.paragraph_format.space_after = Pt(6 if level == 1 else 3)
        _inline(p, text, HEADING_SIZES.get(level, 11), bold=True, font=FONT_HEAD)

    def paragraph(self, text: str):
        _inline(self.doc.add_paragraph(), text)

    def quote(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        _inline(p, text, 10, italic=True)

    def bullet(self, text: str, numbered: bool = False):
        p = self.doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        _inline(p, text)

    def code(self, lines: list[str]):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(12)
        run = p.add_run("\n".join(lines))
        _font(run, FONT_MONO, 9, mono=True)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    def table(self, rows: list[list[str]]):
        if not rows:
            return
        cols = max(len(r) for r in rows)
        t = self.doc.add_table(rows=0, cols=cols)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            for j in range(cols):
                text = row[j] if j < len(row) else ""
                _inline(cells[j].paragraphs[0], text, 9.5, bold=(i == 0), font=FONT_HEAD if i == 0 else FONT_BODY)
        self.doc.add_paragraph()


def _split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def render_markdown(r: Renderer, text: str):
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            block, i = [], i + 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            r.code(block)
            i += 1
            continue

        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not TABLE_SEP.match(lines[i].strip()):
                    rows.append(_split_row(lines[i]))
                i += 1
            r.table(rows)
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            r.heading(len(m.group(1)), m.group(2))
        elif stripped.startswith("> "):
            r.quote(stripped[2:])
        elif re.match(r"^- \[[ x]\]\s+", stripped):
            r.bullet(re.sub(r"^- \[[ x]\]\s+", "", stripped))
        elif stripped.startswith("- "):
            r.bullet(stripped[2:])
        elif re.match(r"^\d+\.\s+", stripped):
            r.bullet(re.sub(r"^\d+\.\s+", "", stripped), numbered=True)
        elif stripped:
            r.paragraph(stripped)
        i += 1


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(title.add_run("智枢·智能体平台 开发设计文档"), FONT_HEAD, 22, True)
    for sub in ("需求说明 / 架构设计 / 数据库设计 / 接口设计", "由 docs/01～04 Markdown 自动生成，以 Markdown 为准"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(p.add_run(sub), FONT_BODY, 11)

    r = Renderer(doc)
    for name in SOURCES:
        doc.add_page_break()
        render_markdown(r, (ROOT / "docs" / name).read_text(encoding="utf-8"))

    doc.save(OUTPUT)
    print("SAVED", OUTPUT.relative_to(ROOT))


if __name__ == "__main__":
    main()
